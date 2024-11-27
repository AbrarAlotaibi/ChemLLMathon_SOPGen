import gradio as gr
from dotenv import load_dotenv
import PyPDF2
from docx import Document
import os
import io
import fitz  # PyMuPDF
from typing import Optional
from docx2txt import process as process_docx
import mimetypes
import tiktoken

from models import (
    get_all_models, 
    get_random_models, 
    process_document,
    count_tokens
)

load_dotenv()
'''
def read_file_content(uploaded_file):
    if uploaded_file is None:
        return None
        
    file_content = None
    
    try:
        # Get file extension and validate
        if isinstance(uploaded_file, tuple):
            filepath = uploaded_file[1]
            file_extension = os.path.splitext(filepath)[1].lower().lstrip('.')
        elif isinstance(uploaded_file, str):
            file_extension = os.path.splitext(uploaded_file)[1].lower().lstrip('.')
        else:
            # Try to detect mime type for binary data
            mime_type = mimetypes.guess_type(uploaded_file.name)[0] if hasattr(uploaded_file, 'name') else None
            if mime_type:
                file_extension = mimetypes.guess_extension(mime_type).lstrip('.')
            else:
                return "Error: Unable to determine file type"

        # Process based on file extension
        if isinstance(uploaded_file, (tuple, str)):
            actual_path = uploaded_file[1] if isinstance(uploaded_file, tuple) else uploaded_file
            
            if file_extension == 'txt':
                with open(actual_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
            
            elif file_extension == 'pdf':
                with open(actual_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    file_content = ""
                    for page in pdf_reader.pages:
                        file_content += page.extract_text()
            
            elif file_extension in ['doc', 'docx']:
                doc = Document(actual_path)
                file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Handle binary data
        else:
            if file_extension == 'pdf':
                pdf_stream = io.BytesIO(uploaded_file.read() if hasattr(uploaded_file, 'read') else uploaded_file)
                pdf_reader = PyPDF2.PdfReader(pdf_stream)
                file_content = ""
                for page in pdf_reader.pages:
                    file_content += page.extract_text()
            
            elif file_extension in ['doc', 'docx']:
                doc_stream = io.BytesIO(uploaded_file.read() if hasattr(uploaded_file, 'read') else uploaded_file)
                doc = Document(doc_stream)
                file_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            elif file_extension == 'txt':
                content = uploaded_file.read() if hasattr(uploaded_file, 'read') else uploaded_file
                try:
                    file_content = content.decode('utf-8') if isinstance(content, bytes) else content
                except UnicodeDecodeError:
                    return "Error: Unable to decode text file. Please ensure it's encoded in UTF-8."

    except Exception as e:
        return f"Error processing file: {str(e)}"
        
    return file_content

# Update the file upload component in your main application
def create_file_upload():
    return gr.File(
        label="Upload Document",
        file_types=["pdf", "doc", "docx", "txt"],
        type="binary",  # Changed from 'filepath' to 'file'
    )
'''    

def read_file_content(uploaded_file) -> Optional[str]:
    """
    Read content from uploaded files (PDF, DOCX, or TXT) using more robust libraries.
    
    Args:
        uploaded_file: The uploaded file object from Gradio
        
    Returns:
        str: The extracted text content or None if processing fails
    """
    if uploaded_file is None:
        return None
        
    try:
        # Handle different input types
        if isinstance(uploaded_file, tuple):
            filepath = uploaded_file[1]
            file_extension = get_file_extension(filepath)
        elif isinstance(uploaded_file, str):
            filepath = uploaded_file
            file_extension = get_file_extension(filepath)
        else:
            # Binary data case
            file_extension = get_file_extension(uploaded_file.name)
            return process_binary_file(uploaded_file, file_extension)
            
        # Process file based on extension
        return process_file_by_path(filepath, file_extension)
            
    except Exception as e:
        return f"Error processing file: {str(e)}"

def get_file_extension(filename: str) -> str:
    """Extract and validate file extension."""
    extension = os.path.splitext(filename)[1].lower().lstrip('.')
    if extension not in ['pdf', 'docx', 'doc', 'txt']:
        raise ValueError(f"Unsupported file type: {extension}")
    return extension

def process_binary_file(file_obj, extension: str) -> Optional[str]:
    """Process file from binary data."""
    try:
        content = file_obj.read()
        if extension == 'pdf':
            with fitz.open(stream=content, filetype="pdf") as doc:
                return "\n".join(page.get_text() for page in doc)
                
        elif extension in ['doc', 'docx']:
            with io.BytesIO(content) as docx_stream:
                return process_docx(docx_stream)
                
        elif extension == 'txt':
            try:
                return content.decode('utf-8')
            except UnicodeDecodeError:
                return content.decode('utf-8', errors='ignore')
                
    except Exception as e:
        raise Exception(f"Error processing binary file: {str(e)}")

def process_file_by_path(filepath: str, extension: str) -> Optional[str]:
    """Process file from filesystem path."""
    try:
        if extension == 'pdf':
            with fitz.open(filepath) as doc:
                return "\n".join(page.get_text() for page in doc)
                
        elif extension in ['doc', 'docx']:
            return process_docx(filepath)
                
        elif extension == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
                
    except Exception as e:
        raise Exception(f"Error processing file at {filepath}: {str(e)}")

# Update the file upload component
def create_file_upload():
    return gr.File(
        label="Upload Document",
        file_types=["pdf", "doc", "docx", "txt"],
        type="binary"
    )
    
def handle_message(
    llms, user_input, uploaded_file, temperature, top_p, max_output_tokens, states1, states2
):
    # Process uploaded file if present
    file_content = None
    if uploaded_file is not None:
        file_content = read_file_content(uploaded_file)
        if file_content:
            file_content = process_document(file_content, user_input)

    history1 = states1.value if states1 else []
    history2 = states2.value if states2 else []
    states = [states1, states2]
    history = [history1, history2]
    
    for hist in history:
        hist.append({"role": "user", "content": user_input})
    
    for updated_history1, updated_history2, updated_states1, updated_states2 in process_responses(
        llms, temperature, top_p, max_output_tokens, history, states, file_content
    ):
        yield updated_history1, updated_history2, updated_states1, updated_states2

def process_responses(llms, temperature, top_p, max_output_tokens, history, states, file_content=None):
    generators = [
        llms[i]["model"](history[i], temperature, top_p, max_output_tokens, file_content)
        for i in range(2)
    ]
    responses = [[], []]
    done = [False, False]

    while not all(done):
        for i in range(2):
            if not done[i]:
                try:
                    response = next(generators[i])
                    if response:
                        responses[i].append(response)
                        history[i][-1] = {
                            "role": "assistant",
                            "content": "".join(responses[i])
                        }
                        states[i] = gr.State(history[i])
                    yield history[0], history[1], states[0], states[1]
                except StopIteration:
                    done[i] = True
    yield history[0], history[1], states[0], states[1]

def activate_chat_buttons():
    regenerate_btn = gr.Button(
        value="🔄  Regenerate", interactive=True, elem_id="regenerate_btn"
    )
    clear_btn = gr.ClearButton(
        elem_id="clear_btn",
        interactive=True,
    )
    return regenerate_btn, clear_btn

def deactivate_chat_buttons():
    regenerate_btn = gr.Button(
        value="🔄  Regenerate", interactive=False, elem_id="regenerate_btn"
    )
    clear_btn = gr.ClearButton(
        elem_id="clear_btn",
        interactive=False,
    )
    return regenerate_btn, clear_btn

with gr.Blocks(
    title="SOPGen",
    theme=gr.themes.Soft(secondary_hue=gr.themes.colors.sky),
) as demo:
    num_sides = 2
    states = [gr.State() for _ in range(num_sides)]
    chatbots = [None] * num_sides
    models = gr.State(get_random_models)
    all_models = get_all_models()
    
    # Get model choices for radio buttons (excluding SOPGen)
    other_models = [model["name"] for model in all_models if model["name"] != "SOPGen"]
    
    gr.HTML("""
        <div style="display: flex; justify-content: space-between; align-items: flex-start; padding: 1rem;">
            <div style="flex: 4;">
                <h1>SOPGen: An LLM SOP Writing Assistant</h1>
                <p>An intelligent documentation specialist customized for chemistry laboratories. This system combines SOP writing expertise with chemistry domain knowledge to help create, review, and optimize laboratory documentation.</p>
                <p>Created as part of Track 3 Problem 2 of the ChemLLM hackathon.</p>
            </div>
            <div style="flex: 1; min-width: 300px; text-align: right;">
                <div style="display: inline-block; text-align: left;">
                    <div style="font-family: Arial, sans-serif; font-size: 32px; font-weight: bold; line-height: 1;">
                        <span style="color: #000;">Chem</span><span style="color: #0099cc;">Indix</span>
                    </div>
                    <div style="font-family: Arial, sans-serif; font-size: 9px; color: #666; margin-top: 5px; letter-spacing: 0.5px;">
                        THE 12TH INTERNATIONAL CONFERENCE
                    </div>
                    <div style="font-family: Arial, sans-serif; font-size: 9px; color: #666; letter-spacing: 0.5px;">
                        AND EXHIBITION ON CHEMISTRY IN INDUSTRY
                    </div>
                </div>
            </div>
        </div>
    """)
    
    with gr.Group(elem_id="share-region-annoy"):
        with gr.Row():
            # Left side - Fixed SOPGen
            with gr.Column():
                chatbots[0] = gr.Chatbot(
                    label="SOPGen",
                    elem_id="chatbot_sopgen",
                    height=550,
                    show_copy_button=True,
                    type="messages"  # Added type
                )
            
            # Right side - Model selection and chatbot
            with gr.Column():
                model_choice = gr.Radio(
                    choices=other_models,
                    label="Select Comparison Model",
                    value=other_models[0],
                    interactive=True
                )
                chatbots[1] = gr.Chatbot(
                    label=model_choice.value,
                    elem_id="chatbot_comparison",
                    height=500,
                    show_copy_button=True,
                    type="messages"  # Added type
                )

    def update_comparison_model(choice):
        all_models = get_all_models()
        selected_model = next(model for model in all_models if model["name"] == choice)
        sopgen = next(model for model in all_models if model["name"] == "SOPGen")
        models.value = [sopgen, selected_model]
        return gr.update(label=choice)

    model_choice.change(
        fn=update_comparison_model,
        inputs=[model_choice],
        outputs=[chatbots[1]]
    )

    
        # Update the main block where you define the file upload
    with gr.Row():
        with gr.Column(scale=4):
            textbox = gr.Textbox(
                show_label=False,
                placeholder="Enter your query and press ENTER",
                elem_id="input_box",
            )
        with gr.Column(scale=1):
            file_upload = create_file_upload()
        with gr.Column(scale=1):  # Changed from 0.5 to 1 to fix the warning
            send_btn = gr.Button(value="Upload", variant="primary")
            
        with gr.Row() as button_row:
            clear_btn = gr.ClearButton(
                value="🎲 New Round",
                elem_id="clear_btn",
                interactive=False,
                components=chatbots + states + [file_upload],
            )
            regenerate_btn = gr.Button(
                value="🔄 Regenerate", interactive=False, elem_id="regenerate_btn"
            )

    with gr.Row():
        examples = gr.Examples(
            [
                "Write an SOP for using HPLC to analyze pharmaceutical samples. Include calibration and sample preparation steps.",
                "Create a comprehensive SOP for handling and disposal of hazardous chemical waste in a research laboratory.",
                "Generate an SOP for the pH meter calibration process using standard buffer solutions and include quality control steps."
            ],
            inputs=[textbox],
            label="Example inputs"
        )

    with gr.Accordion("Parameters", open=False) as parameter_row:
        temperature = gr.Slider(
            minimum=0.0,
            maximum=2.0,
            value=0.0,
            step=0.1,
            interactive=True,
            label="Temperature",
        )
        top_p = gr.Slider(
            minimum=0.0,
            maximum=1.0,
            value=1.0,
            step=0.1,
            interactive=True,
            label="Top P",
        )
        max_output_tokens = gr.Slider(
            minimum=16,
            maximum=4096,
            value=1024,
            step=64,
            interactive=True,
            label="Max output tokens",
        )

    textbox.submit(
        handle_message,
        inputs=[
            models,
            textbox,
            file_upload,
            temperature,
            top_p,
            max_output_tokens,
            states[0],
            states[1],
        ],
        outputs=[chatbots[0], chatbots[1], states[0], states[1]],
    ).then(
        activate_chat_buttons,
        inputs=[],
        outputs=[regenerate_btn, clear_btn],
    )

    send_btn.click(
        handle_message,
        inputs=[
            models,
            textbox,
            file_upload,
            temperature,
            top_p,
            max_output_tokens,
            states[0],
            states[1],
        ],
        outputs=[chatbots[0], chatbots[1], states[0], states[1]],
    ).then(
        activate_chat_buttons,
        inputs=[],
        outputs=[regenerate_btn, clear_btn],
    )

    regenerate_btn.click(
        handle_message,
        inputs=[
            models,
            textbox,
            file_upload,
            temperature,
            top_p,
            max_output_tokens,
            states[0],
            states[1],
        ],
        outputs=[chatbots[0], chatbots[1], states[0], states[1]],
    )

    clear_btn.click(
        deactivate_chat_buttons,
        inputs=[],
        outputs=[regenerate_btn, clear_btn],
    ).then(lambda: get_random_models(), inputs=None, outputs=[models])

if __name__ == "__main__":
    demo.queue(default_concurrency_limit=10)
    demo.launch()



  